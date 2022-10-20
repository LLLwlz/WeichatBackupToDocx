import hashlib
import io
import os
import re
import shutil
import time
import itertools
import xml.etree.ElementTree as ET
from urllib.request import urlretrieve

import docx
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
import requests
from docx import shared
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_COLOR_INDEX
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Pt
from pyquery import PyQuery
from pysqlcipher3 import dbapi2 as sqlite
from termcolor import colored

# TYPE模式的宏定义
TYPE_MSG = 1
TYPE_IMG = 3
TYPE_SPEAK = 34
TYPE_NAME_CARD = 42
TYPE_VIDEO_FILE = 43
TYPE_BIG_EMOJI = 47
TYPE_LINK = 49  # 链接共享或来自网络的文件, see https://github.com/ppwwyyxx/wechat-dump/issues/52
TYPE_VOIP = 50  # 语音通话
TYPE_SYSTEM = 10000  # 撤回/转账接收/拒收消息
TYPE_CUSTOM_EMOJI = 1048625
TYPE_WITHDRAW_MSG = 268445456  # 撤回的消息
TYPE_MONEY_TRANSFER = 419430449  # 微信转账
TYPE_LUCKY_MONEY = 436207665  # 发红包
TYPE_ANSWER_MSG = 822083633  # 回复消息
TYPE_SHOT = 922746929  # 拍一拍
TYPE_FILE = 1090519089  # 发送文件
TYPE_APP_MSG = 16777265

MSG_MAX = 20

def remove_control_chars(s):
    """
    去除可能存在的不可见控制字符

    :param s:
    :return:
    """
    control_chars = ''.join(map(chr, itertools.chain(range(0x00, 0x20), range(0x7f, 0xa0))))
    control_char_re = re.compile('[%s]' % re.escape(control_chars))
    return control_char_re.sub('', s)

def connect_wcdb():
    """
    连接数据库

    :return:
    """
    msg = sqlite.connect('EnMicroMsg-decrypted.db')
    msg_c = msg.cursor()

    file = sqlite.connect('WxFileIndex-decrypted.db')
    file_c = file.cursor()

    return msg_c, file_c, msg, file


def close_wcdb(msg_c, file_c, msg, file):
    """
    关闭数据库

    :param msg_c:
    :param file_c:
    :param msg:
    :param file:
    """
    msg.commit()
    file.commit()
    msg_c.close()
    file_c.close()


def get_md5(text):
    """
    将wxid使用MD5编码加密, 加密结果是用户头像路径

    :param text: 需要加密的字符串
    :return:
    """
    m = hashlib.md5()
    # 参数必须是byte类型，否则报Unicode-objects must be encoded before hashing错误
    m.update(bytes(text.encode('utf-8')))
    return m.hexdigest()


def get_avator_path(wxid):
    """
    获取头像文件完整路径

    :param wxid:
    :return:
    """
    avatar = get_md5(wxid)
    # print(avatar)
    avatar_path = r"./avatar/"
    path = avatar_path + avatar[:2] + '/' + avatar[2:4]
    for root, dirs, files in os.walk(path):
        for file in files:
            if avatar in file:
                avatar = file
                break
    return path + '/' + avatar


def get_emoji_path(msg, imgPath):
    """
    获取表情包路径

    :param msg: content
    :param imgPath:
    :return: ret, 路径
    """
    try:
        Path = r'./emoji'
        for root, dirs, files in os.walk(Path):
            for file in files:
                if imgPath + '_cover' in file:
                    emoji_path = root + '/' + file
                    return 1, emoji_path
        for root, dirs, files in os.walk(Path):
            for file in files:
                if imgPath in file:
                    emoji_path = root + '/' + file
                    return 1, emoji_path
        # path = f'.//emoji//{imgPath}'
        # is_Exist = os.path.exists(path)
        # if not is_Exist:
        '''表情包不存在，则下载表情包到emoji文件夹中'''
        ret = download_emoji(msg, imgPath)
        if ret == 0:
            return 0, 0
        emoji_path = f'./emoji/{imgPath}'
        return 1, emoji_path
        # image(doc, isSend, Type=47, content=content, imgPath=imgPath)
    except Exception:
        return 0, 0
        # print("can't find emoji!")


def download_emoji(content, img_path):
    """
    下载emoji文件

    :param content:
    :param img_path:
    :return: ret
    """
    try:
        url = content.split('cdnurl = "')[1].split('"')[0]
        url = ':'.join(url.split('*#*'))
        if 'amp;' in url:
            url = ''.join(url.split('amp;'))
            # print('emoji downloading!!!')
        # print(url)
        # print('1')
        resp = requests.get(url)
        with open(f'./emoji/{img_path}', 'wb') as f:
            f.write(resp.content)
        return 1
    except Exception:
        # print("emoji download error")
        return 0


'''获得爱人的聊天数据，返回字典'''


def get_love_msg(conRemark, timeStart, timeEnd):
    """
    获得爱人的聊天数据

    :param conRemark: 备注
    :param timeStart: 开始时间
    :param timeEnd: 结束时间
    :return: 聊天数据字典
    """
    select_love_msg = 'SELECT datetime(subStr(cast(msg.createTime as text),1,10),"unixepoch","localtime") AS theTime,' \
                      'CASE msg.isSend ' \
                      'WHEN 0 THEN name.conRemark ' \
                      'WHEN 1 THEN "我" ' \
                      'END AS person,' \
                      'msg.isSend AS isSend,' \
                      'msg.content AS message,' \
                      'msg.type AS type,' \
                      'msg.imgPath AS imgPath,' \
                      'msg.status AS status,' \
                      'msg.msgId AS msgId,' \
                      'msg.lvbuffer AS buffer ' \
                      'FROM message msg INNER JOIN rcontact name ON msg.talker = name.username ' \
                      'WHERE name.conRemark = \'' + conRemark + '\' ' \
                                                                'AND theTime < \'' + timeEnd + '\' ' \
                                                                                               'AND theTime >= \'' + timeStart + '\' ' \
                                                                                                                                 'ORDER BY msg.createTime;'
    # 'AND msg.type = 49 ' \
    msg_cur, file_cur, msg_con, file_con = connect_wcdb()
    msg_cur.execute(select_love_msg)
    desc = msg_cur.description
    special_love_dict = [dict(zip([col[0] for col in desc], row)) for row in msg_cur.fetchall()]
    close_wcdb(msg_cur, file_cur, msg_con, file_con)

    return special_love_dict


def get_love_wxid(conRemark):
    """
    获取指定备注名的wxid

    :param conRemark:
    :return: wxid
    """
    select_love_wxid = 'SELECT username ' \
                       'FROM rcontact ' \
                       'WHERE conRemark = \'' + conRemark + '\';'

    msg_cur, file_cur, msg_con, file_con = connect_wcdb()
    msg_cur.execute(select_love_wxid)
    wxid = msg_cur.fetchall()[0][0]
    close_wcdb(msg_cur, file_cur, msg_con, file_con)
    # print(wxid)
    return wxid


def get_self_wxid():
    """
    获取自己的wxid

    :return: wxid
    """
    select_self_wxid = 'SELECT value ' \
                       'FROM userinfo ' \
                       'WHERE value like "wxid#_%" ESCAPE "#";'
    msg_cur, file_cur, msg_con, file_con = connect_wcdb()
    msg_cur.execute(select_self_wxid)
    wxid = msg_cur.fetchall()[0][0]
    close_wcdb(msg_cur, file_cur, msg_con, file_con)
    # print(wxid)
    return wxid


def IS_3_min(last_time, now_time):
    """
    判断两次聊天时间是不是大于三分钟。若大于三分钟则显示时间；否则不显示

    :param last_time:
    :param now_time:
    :return: ret(是否显示)
    """
    last = time.strptime(last_time, "%Y-%m-%d %H:%M:%S")
    now = time.strptime(now_time, "%Y-%m-%d %H:%M:%S")
    '''两次聊天记录时间差，单位是秒'''
    love_sub = (now.tm_sec - last.tm_sec) + (now.tm_min - last.tm_min) * 60 + (now.tm_hour - last.tm_hour) * 60 * 60 + (
            now.tm_yday - last.tm_yday) * 24 * 60 * 60 + (now.tm_year - last.tm_year) * 366 * 24 * 60 * 60
    # print(sub)
    return love_sub >= 180


def IS_8_hour(last_time, now_time):
    """
    判断上次打印时间是不是大于八小时。若大于八小时则打印时间；否则不打印

    :param last_time:
    :param now_time:
    :return: ret(是否打印)
    """
    last = time.strptime(last_time, "%Y-%m-%d %H:%M:%S")
    now = time.strptime(now_time, "%Y-%m-%d %H:%M:%S")
    '''两次的打印时间差，单位是秒'''
    love_sub = (now.tm_sec - last.tm_sec) + (now.tm_min - last.tm_min) * 60 + (now.tm_hour - last.tm_hour) * 60 * 60 + (
            now.tm_yday - last.tm_yday) * 24 * 60 * 60 + (now.tm_year - last.tm_year) * 366 * 24 * 60 * 60
    # print(sub)
    return love_sub >= 28800


def text_love_time(doc, love_time):
    """
    输出时间

    :param doc:
    :param love_time:
    """
    love_paragraph = doc.add_paragraph("")
    love_run = love_paragraph.add_run(love_time)
    love_run.font.name = 'Times New Roman'
    love_run.font.size = Pt(11)
    love_paragraph.paragraph_format.space_before = Pt(10)
    love_paragraph.paragraph_format.space_after = Pt(10)
    love_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def create_love_table(doc, isSend, person):
    """
    创建一个1*2表格：isSend = 1 (0,0)存聊天内容，(0,1)存头像；isSend = 0 (0,0)存头像，(0,1)存聊天内容

    :param doc:
    :param isSend:
    :param person:
    :return: 聊天内容的坐标
    """
    love_table = doc.add_table(rows=1, cols=2, style='Normal Table')
    love_table.cell(0, 1).height = shared.Inches(0.5)
    love_table.cell(0, 0).height = shared.Inches(0.5)
    if isSend:
        self_wxid = get_self_wxid()
        self_avator = get_avator_path(self_wxid)
        img_self = open(self_avator, 'rb')
        '''表格右对齐'''
        love_table.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        avatar = love_table.cell(0, 1).paragraphs[0]
        '''插入头像，设置头像宽度'''
        avatar.add_run().add_picture(img_self, width=shared.Inches(0.5))
        '''设置单元格宽度跟头像一致'''
        love_table.cell(0, 1).width = shared.Inches(0.5)
        content_cell = love_table.cell(0, 0)
        '''聊天内容右对齐'''
        # content_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        img_self.close()
    else:
        love_wxid = get_love_wxid(person)
        love_avator = get_avator_path(love_wxid)
        img_love = open(love_avator, 'rb')
        avatar = love_table.cell(0, 0).paragraphs[0]
        '''插入头像，设置头像宽度'''
        avatar.add_run().add_picture(img_love, width=shared.Inches(0.5))
        '''设置单元格宽度跟头像一致'''
        love_table.cell(0, 0).width = shared.Inches(0.5)
        content_cell = love_table.cell(0, 1)
        img_love.close()
    '''聊天内容垂直居中对齐'''
    content_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    avatar.paragraph_format.space_before = Pt(0)
    avatar.paragraph_format.space_after = Pt(0)
    avatar.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # content_cell.width = shared.Inches(8)
    return content_cell


def text_love(doc, isSend, person, message, status):
    """
    将文字聊天记录写入文件

    :param doc: 
    :param isSend: 是谁发送的信息
    :param person: 
    :param message: 
    :param status: 状态码
    """
    if status == 5:
        message += '（未发出） '
    content_cell = create_love_table(doc, isSend, person)
    try:
        content_run = content_cell.paragraphs[0].add_run(message)
    except:
        content_run = content_cell.paragraphs[0].add_run(remove_control_chars(message))
    content_cell.paragraphs[0].font_size = shared.Inches(0.5)
    content_run.font.size = Pt(12)
    content_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    content_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    '''当且仅当是自己发送的且发送的字符较短时：左对齐'''
    if len(message) < MSG_MAX and isSend:
        content_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


def image_love(doc, isSend, person, imgPath):
    """
    插入聊天图片：isSend = 1 只有缩略图，isSend = 0 有原图 

    :param doc: 
    :param isSend: 
    :param person: 
    :param imgPath: 
    """
    content_cell = create_love_table(doc, isSend, person)
    content_run = content_cell.paragraphs[0].add_run()
    content_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    content_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    imgPath = imgPath.split('//th_')[-1]
    Path = f'.//image2//{imgPath[:2]}//{imgPath[2:4]}'
    for root, dirs, files in os.walk(Path):
        for file in files:
            if isSend:
                if imgPath + 'hd' in file:
                    if '_hevc' in file:
                        continue
                    imgPath = file
                    break
            if imgPath in file:
                if '_hevc' in file:
                    continue
                if 'jpg' in file:
                    continue
                imgPath = file
                break
    try:
        img_love = open(f'{Path}/{imgPath}', 'rb')
        '''插入图片，设置单元格高度跟图片一致'''
        content_run.add_picture(img_love, height=shared.Inches(2))
        content_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        content_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        if isSend:
            content_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        # content_cell.width = shared.Inches(0.5)
        '''注意关闭图片，否则不能移植'''
        img_love.close()
        # content_cell.paragraphs[0].add_run('您的图片已过期或被错误删除')
        # doc.add_paragraph()
    except Exception:
        # print(f'{Path}/{imgPath}')
        # print('您的图片已过期或被错误删除')
        content_run = content_cell.paragraphs[0].add_run('您的图片已过期或被错误删除')
        content_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        content_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        content_run.font.size = Pt(11)
        content_run.font.color.rgb = shared.RGBColor(191, 191, 191)
        content_run.font.strike = True
        if isSend:
            content_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        # print("Error!image")


def emoji_love(doc, isSend, person, message, imgPath):
    """
    插入聊天表情包
    
    :param doc: 
    :param isSend: 
    :param person: 
    :param message: 
    :param imgPath: 
    """
    ret, emoji_path = get_emoji_path(message, imgPath)
    # print(emoji_path)
    content_cell = create_love_table(doc, isSend, person)
    content_run = content_cell.paragraphs[0].add_run()
    if ret:
        emoji_img_love = open(emoji_path, 'rb')
        try:
            content_run.add_picture(emoji_img_love, height=shared.Inches(2))
            content_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            content_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        except Exception:
            # print(imgPath)
            # print('不能正确打开表情')
            content_run = content_cell.paragraphs[0].add_run('不能正确打开表情')
            content_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            content_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            content_run.font.size = Pt(11)
            content_run.font.color.rgb = shared.RGBColor(191, 191, 191)
            content_run.font.strike = True
        if isSend:
            content_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        # content_cell.width = shared.Inches(0.5)
        emoji_img_love.close()
    else:
        # print(imgPath)
        # print('不能正确找到表情')
        content_run = content_cell.paragraphs[0].add_run('不能正确找到表情')
        content_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        content_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        content_run.font.size = Pt(11)
        content_run.font.color.rgb = shared.RGBColor(191, 191, 191)
        content_run.font.strike = True
        if isSend:
            content_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


def voice_love(doc, isSend, person, imgPath, tmp_voice_dic):
    """
    插入微信语音，利用超链接的形式
    
    :param tmp_voice_dic:
    :param doc:
    :param isSend: 
    :param person: 
    :param imgPath: 
    :return: 
    """
    voice = get_md5(imgPath)
    voice_file = './voice2/' + voice[:2] + '/' + voice[2:4] + '/msg_' + imgPath + '.amr'
    content_cell = create_love_table(doc, isSend, person)
    # print(voice_file)
    try:
        '''复制一份到新的文件夹中以免文件过大'''
        shutil.copyfile(voice_file, tmp_voice_dic + '/' + 'msg_' + imgPath + '.amr')
    except Exception:
        # print(voice_file)
        # print('您的语音已过期或被错误删除')
        content_run = content_cell.paragraphs[0].add_run('您的语音已过期或被错误删除')
        content_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        content_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        content_run.font.size = Pt(11)
        content_run.font.color.rgb = shared.RGBColor(191, 191, 191)
        content_run.font.strike = True
        if isSend:
            content_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        return

    add_hyperlink(content_cell, person + '的语音', 'voice_love/' + 'msg_' + imgPath + '.mp3')
    content_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    content_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    if isSend:
        content_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


def content_xml_ready(content):
    """
    利用content构建xml

    :param content:
    :return: 预处理好的xml
    """
    '''删除 xml 标头以避免可能产生的错误'''
    header = re.compile(r'<\?.*\?>')
    msg = header.sub("", content)
    return msg


def reply_love(doc, isSend, person, message, status):
    """
    添加回复信息

    :param doc:
    :param isSend:
    :param person:
    :param message:
    :param status:
    """
    # print(message)
    flag_emoji = 0
    '''获取回复信息'''
    answer = re.compile(r"<title>(?P<title>(.*?))</title>")
    answer_love = answer.search(message).groupdict()['title']
    # print(answer_love)
    '''获取回复的人'''
    name = re.compile(r"<displayname>(?P<displayname>(.*?))</displayname>")
    name_love = name.search(message).groupdict()['displayname']
    # print(name_love)
    '''获取要回复的那条信息'''
    last = re.compile(r"\n?title&gt;(?P<content>(.*?))\n?&lt;/title&gt")
    if not last.search(message):
        if isSend == 0:
            '''匹配对方的回复'''
            last = re.compile(r"<content>(?P<content>(.*?))</content>")
        else:
            '''匹配自己的回复'''
            last = re.compile(r"</msgsource>\n?<content>(?P<content>(.*?))\n?</content>")

    try:
        '''试错'''
        last_love = last.search(message).groupdict()['content']
        # print(last_love)
    except Exception:
        try:
            '''试错'''
            last_love = last.search(message).groupdict()['content']
        except Exception:
            '''试错'''
            last = re.compile(r"\n?<content>(?P<content>(.*?))\n?</content>")
            '''试错'''
            if last.search(message):
                last_love = last.search(message).groupdict()['content']
            else:
                # print(message)
                last_love = '[图片]'
            # TODO: 解决图片/语音引用问题
    if status == 5:
        message += '（未发出） '
    if 'xml' in last_love:
        if len(last_love) > 2 * MSG_MAX:
            last_love = '[图片]'
    if '&quot;http*#*//' in last_love:
        if len(last_love) > 2 * MSG_MAX:
            url = re.compile(r"&quot;http\*#\*//(?P<content>(.*?))&quot;&#x20;")
            try:
                love_url = url.search(last_love).groupdict()['content']
                # print(message)
                # print(love_url)
                flag_emoji = 1
            except Exception:
                flag_emoji = 0
    if 'wxid' in last_love:
        last_love = '[动画表情]'
    content_cell = create_love_table(doc, isSend, person)
    content_run = content_cell.paragraphs[0].add_run(answer_love)
    content_run.font.size = Pt(12)
    content_cell.paragraphs[0].font_size = shared.Inches(0.5)
    content_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    content_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    if len(answer_love) < MSG_MAX and isSend:
        p = content_cell.paragraphs[0]
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    reply_p = content_cell.add_paragraph()
    if flag_emoji:
        try:
            urlretrieve('http://' + love_url, './tmp/' + 'new_img' + '.png')
            img_emoji = open('./tmp/' + 'new_img' + '.png', 'rb')
            run = content_cell.paragraphs[1].add_run(name_love + ':')
            run.add_picture(img_emoji, height=shared.Inches(2))
            run.font.color.rgb = shared.RGBColor(121, 121, 121)
            run.font_size = shared.Inches(0.3)
            run.font.size = Pt(11)
            run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
            content_cell.paragraphs[1].paragraph_format.space_before = Pt(0)
            content_cell.paragraphs[1].paragraph_format.space_after = Pt(0)
            if isSend:
                reply_p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            return
        except Exception:
            last_love = '[动画表情]'

    run = content_cell.paragraphs[1].add_run(name_love + ':' + last_love)
    '''设置被回复内容格式'''
    run.font.color.rgb = shared.RGBColor(121, 121, 121)
    run.font_size = shared.Inches(0.3)
    run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
    content_cell.paragraphs[1].paragraph_format.space_before = Pt(0)
    content_cell.paragraphs[1].paragraph_format.space_after = Pt(0)
    run.font.size = Pt(11)
    # TODO: 优化引用格式
    if len(name_love + ':' + last_love) < MSG_MAX and isSend:
        reply_p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


def add_hyperlink(content_cell, text, url):
    """
    添加超链接

    :param content_cell:
    :param text: 文本
    :param url: 链接
    """
    # paragraph = doc.add_paragraph('')
    paragraph = content_cell.paragraphs[0]

    '''关联超链接'''
    r_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    run = paragraph.add_run(text)
    run.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    run.font.underline = True
    run.font.size = Pt(12)
    hyperlink.append(run._r)
    paragraph._element.append(hyperlink)
    return hyperlink


def retract_message_love(doc, message):
    """
    添加撤回信息提示

    :param doc:
    :param message:
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(message)
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run.font.size = Pt(10.5)
    run.font.color.rgb = shared.RGBColor(191, 191, 191)


def mkdir(path):
    """
    新建文集所需文件夹

    :param path:
    :return: 是否被创建
    """
    path = path.strip()
    path = path.rstrip("\\")
    if os.path.exists(path):
        return False
    os.makedirs(path)
    return True


def namecard_love(doc, isSend, person, message):
    """
    添加分享名片信息

    :param doc:
    :param isSend:
    :param person:
    :param message:
    """
    content_cell = create_love_table(doc, isSend, person)
    content_run = content_cell.paragraphs[0].add_run()
    pq_love = PyQuery(content_xml_ready(message), parser='xml')
    msg_love = pq_love('msg').attr
    nickname = msg_love['nickname']
    if not nickname:
        nickname = msg_love['alias']
    if not nickname:
        nickname = ""
    download_url = msg_love['brandIconUrl']
    if not download_url:
        user_name = msg_love['username']
    try:
        urlretrieve(download_url, './tmp/' + nickname + '.png')
    except Exception:
        try:
            love_avator = get_avator_path(user_name)
        except Exception:
            print("Can't deal name_card\n" + message)
    run0 = content_cell.add_paragraph().add_run('分享卡片')
    run0.font.size = Pt(12)
    try:
        img_log = open('./tmp/' + nickname + '.png', 'rb')
        content_run.add_picture(img_log, height=shared.Inches(2))
    except Exception:
        try:
            img_avator = open(love_avator, 'rb')
            content_run.add_picture(img_avator, height=shared.Inches(2))
        except Exception:
            run1 = content_cell.paragraphs[0].add_run('[图片]')
            run1.font.size = Pt(12)
    run2 = content_cell.add_paragraph().add_run(nickname)
    run2.font.size = Pt(12)

    content_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    content_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    content_cell.paragraphs[1].paragraph_format.space_before = Pt(0)
    content_cell.paragraphs[1].paragraph_format.space_after = Pt(0)
    content_cell.paragraphs[2].paragraph_format.space_before = Pt(0)
    content_cell.paragraphs[2].paragraph_format.space_after = Pt(0)

    if isSend:
        content_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        content_cell.paragraphs[1].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        content_cell.paragraphs[2].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


def voip_love(doc, isSend, person, buffer):
    """
    添加语音通话结束标志

    :param doc:
    :param isSend:
    :param person:
    :param buffer:
    """
    content_cell = create_love_table(doc, isSend, person)
    str_tmp = str(buffer, 'utf-8')
    if str_tmp.find('聊天时长') != -1:
        str_love = str_tmp[str_tmp.find('聊天时长'):str_tmp.find(':') + 3]
    elif str_tmp.find('通话时长') != -1:
        str_love = str_tmp[str_tmp.find('通话时长'):str_tmp.find(':') + 3]
    elif str_tmp.find('呼叫失败') != -1:
        str_love = u'呼叫失败'
    elif str_tmp.find('通话中断') != -1:
        try:
            str_love = str_tmp[str_tmp.find('通话中断'):str_tmp.find(':') + 3]
        except Exception:
            str_love = u'通话中断'
    elif str_tmp.find('对方无应答') != -1:
        str_love = u'对方无应答'
    elif str_tmp.find('已拒绝') != -1:
        str_love = u'对方已拒绝'
    elif str_tmp.find('已取消') != -1:
        str_love = u'已取消'
    elif str_tmp.find('忙线') != -1:
        str_love = u'忙线未接听'
    elif str_tmp.find('连接失败') != -1:
        str_love = u'连接失败'
    elif str_tmp.find('呼叫失败') != -1:
        str_love = u'呼叫失败'
    else:
        # print(str_tmp)
        str_love = '语音通话'
    # print(str_love)
    content_run = content_cell.paragraphs[0].add_run(str_love)
    content_cell.paragraphs[0].font_size = shared.Inches(0.5)
    content_run.font.size = Pt(12)
    content_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    content_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    if isSend:
        content_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


def file_love(doc, isSend, person, message, file_love_dic):
    """
    插入发送的文件，利用超链接的形式

    :param doc:
    :param isSend:
    :param person:
    :param message:
    :param file_love_dic:
    :return:
    """
    file = re.compile(r"<title>(.*?)<")
    love_file = file.search(message).group()
    love_filename = love_file.lstrip('<title>').rstrip('<')
    content_cell = create_love_table(doc, isSend, person)
    try:
        '''复制一份到新的文件夹中以免文件过大'''
        shutil.copyfile('./Download/' + love_filename, file_love_dic + '/' + love_filename)
    except Exception:
        # print(love_filename)
        # print('您的文件已过期或被错误删除')
        content_run = content_cell.paragraphs[0].add_run(love_filename + '\n(您的文件已过期或被错误删除)')
        content_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        content_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        content_run.font.size = Pt(11)
        content_run.font.color.rgb = shared.RGBColor(191, 191, 191)
        content_run.font.strike = True
        if isSend:
            content_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            # content_cell.paragraphs[1].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        return

    add_hyperlink(content_cell, love_filename, 'file_love/' + love_filename)
    content_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    content_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    if isSend:
        content_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    # print(love_filename)


def shot_love(doc, person, message):
    """
    添加拍一拍信息

    :param doc:
    :param person:
    :param message:
    """
    pattern = re.compile(r"<template><!\[CDATA\[(?P<it>(.*?))]]></template>")
    result = pattern.search(message).groupdict()['it']
    fromusername = '${fromusername@textstatusicon}'
    pattedusername = '${pattedusername@textstatusicon}'
    '''我拍别人'''
    if result[0] == u'我':
        love_pat = result
    else:
        '''处理多余的引号'''
        result = result.split('""') if '""' in result else result.split('"')
        love_pat = '"' + person + '"' + result[2]
    love_pat = ''.join(love_pat.split(fromusername))
    love_pat = ''.join(love_pat.split(pattedusername))
    love_pat = ''.join(love_pat.split(' '))
    # print(love_pat)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(love_pat)
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    '''设置拍一拍文字格式'''
    run.font.color.rgb = shared.RGBColor(121, 121, 121)
    run.font_size = shared.Inches(0.3)
    run.font.bold = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    # run.font.highlight_color=WD_COLOR_INDEX.GRAY_25


def system_love(doc, person, message):
    """
    添加模式提示：领取红包/撤回/拒收/转账提示

    :param doc:
    :param person:
    :param message:
    """
    if '撤回' in message:
        love_content = message
    elif '领取' in message:
        love_content = person + '领取了你的红包'
    elif '拒收' in message:
        love_content = message
    elif '将于2小时后过期' in message:
        love_content = '你有一笔待接收的转账，将于2小时后过期'
    elif '收款方24小时内未接收你的' in message:
        love_content = '你有一笔待接收的转账，将于2小时后过期'
    elif '24小时内未接收，已过期' in message:
        love_content = '你有一笔待接收的转账，将于2小时后过期'
    elif '你有一笔待接收的' in message:
        love_content = '你有一笔待接收的转账'
    else:
        print("Can't deal system\n" + message)
        love_content = message
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(love_content)
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run.font.size = Pt(10.5)
    run.font.color.rgb = shared.RGBColor(191, 191, 191)
    # run.font.highlight_color=WD_COLOR_INDEX.GRAY_25


def video_love(doc, isSend, person, imgPath, video_love_dic):
    """
    插入视频文件，包括封面图和视频全部

    :param video_love_dic:
    :param doc:
    :param isSend:
    :param person:
    :param imgPath:
    """
    content_cell = create_love_table(doc, isSend, person)
    content_run = content_cell.paragraphs[0].add_run()
    Path = './video'
    for root, dirs, files in os.walk(Path):
        for file in files:
            if imgPath in file:
                if '.jpg' in file:
                    love_image = file
                if '.mp4' in file:
                    love_video = file
    try:
        content_run.add_picture(f'{Path}/{love_image}', height=shared.Inches(2))
        content_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        content_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        shutil.copyfile(Path + '/' + love_video, video_love_dic + '/' + love_video)
        add_hyperlink(content_cell, '\n播放视频', 'video_love/' + love_video)

    except Exception:
        # print(f'{Path}/{imgPath}')
        # print('您的视频已过期或被错误删除')
        content_run = content_cell.paragraphs[0].add_run('您的视频已过期或被错误删除')
        content_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        content_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        content_run.font.size = Pt(11)
        content_run.font.color.rgb = shared.RGBColor(191, 191, 191)
        content_run.font.strike = True

    if isSend:
        content_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


def link_love(doc, isSend, person, message, file_love_dic):
    """
    添加微信分享信息：聊天记录/网络文件/外部链接

    :param file_love_dic:
    :param doc:
    :param isSend:
    :param person:
    :param message:
    :return:
    """
    content_cell = create_love_table(doc, isSend, person)
    # content_run = content_cell.paragraphs[0].add_run()
    pq_love = PyQuery(content_xml_ready(message))
    url_love = pq_love('url').text()
    title_love = pq_love('title').text().split(' null')[0]
    if '聊天记录' in title_love:
        des_love = pq_love('des').text()
        content_run = content_cell.paragraphs[0].add_run(title_love)
        content_cell.paragraphs[0].font_size = shared.Inches(0.5)
        content_run.font.size = Pt(12)
        content_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        content_cell.paragraphs[0].paragraph_format.space_after = Pt(5)
        # des_love = ''.join(des_love.split('\n'))
        # content_cell.add_paragraph('\n'.join(des_love.split(' ')))
        content_cell.add_paragraph()
        # des_love = '\n'.join(des_love.split('\r'))
        # print(des_love)
        content_run = content_cell.paragraphs[1].add_run(des_love)
        content_cell.paragraphs[1].font_size = shared.Inches(0.5)
        content_run.font.size = Pt(12)
        content_cell.paragraphs[1].paragraph_format.space_before = Pt(0)
        content_cell.paragraphs[1].paragraph_format.space_after = Pt(0)
        return
    elif not url_love:
        if pq_love('des').text() != '':
            line_love = title_love + '(' + pq_love('des').text() + ')'
        else:
            line_love = title_love
        try:
            shutil.copyfile('./Download/' + title_love, file_love_dic + '/' + title_love)
            add_hyperlink(content_cell, line_love, 'file_love/' + title_love)
            content_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            content_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            # print('超链接')
        except Exception:
            content_run = content_cell.paragraphs[0].add_run(line_love)
            content_cell.paragraphs[0].font_size = shared.Inches(0.5)
            content_run.font.size = Pt(12)
            content_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            content_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        if isSend:
            content_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        return
    elif url_love:
        add_hyperlink(content_cell, title_love, url_love)
        content_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        content_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        if isSend:
            content_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        # print("URL:{}".format(url_love))
        # print(content_xml_ready(message))
        return
    else:
        print("Can't deal link\n" + content_xml_ready(message))
        content_run = content_cell.paragraphs[0].add_run(content_xml_ready(message))
        content_cell.paragraphs[0].font_size = shared.Inches(0.5)
        content_run.font.size = Pt(12)
        content_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        content_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        if isSend:
            content_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        # print(content_xml_ready(message))


def custom_emoji_love(doc, isSend, person, message):
    """
    添加用户表情包

    :param doc:
    :param isSend:
    :param person:
    :param message:
    """
    if 'emoticonmd5' in message:
        pq_love = PyQuery(message)
        imgPath = pq_love('emoticonmd5').text()
        # print(imgPath)
        emoji_love(doc, isSend, person, message, imgPath)
    else:
        print('Custom emoji deal wrong')


def money_transfer_love(doc, isSend, person, message):
    """
    添加转账信息

    :param doc:
    :param isSend:
    :param person:
    :param message:
    """
    content_cell = create_love_table(doc, isSend, person)
    data_to_parse = io.BytesIO(message.encode('utf-8'))
    love_money = "[微信转账]"
    try:
        for event, elem in ET.iterparse(data_to_parse, events=('end',)):
            if elem.tag == 'des':
                title = elem.text
                love_money = "[微信转账]\n{}".format(title)
                break
    except Exception:
        love_money = "[微信转账]"
    content_run = content_cell.paragraphs[0].add_run(love_money)
    content_cell.paragraphs[0].font_size = shared.Inches(0.5)
    content_run.font.size = Pt(12)
    content_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    content_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    if isSend:
        content_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


def luck_money_love(doc, isSend, person, message):
    """
    添加微信红包信息

    :param doc:
    :param isSend:
    :param person:
    :param message:
    """
    content_cell = create_love_table(doc, isSend, person)
    data_to_parse = io.BytesIO(message.encode('utf-8'))
    love_lucky_money = "[微信红包]"
    try:
        for event, elem in ET.iterparse(data_to_parse, events=('end',)):
            if elem.tag == 'sendertitle':
                title = elem.text
                love_lucky_money = "[微信红包]\n{}".format(title)
                break
    except Exception:
        love_lucky_money = "[微信红包]"
    content_run = content_cell.paragraphs[0].add_run(love_lucky_money)
    content_cell.paragraphs[0].font_size = shared.Inches(0.5)
    content_run.font.size = Pt(12)
    content_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    content_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    if isSend:
        content_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


def app_msg_love(doc, isSend, person, message):
    """
    添加APP到微信信息

    :param doc:
    :param isSend:
    :param person:
    :param message:
    """
    content_cell = create_love_table(doc, isSend, person)
    pq_love = PyQuery(content_xml_ready(message))
    des_love = pq_love('des').text()
    title_love = pq_love('title').text()
    content_run = content_cell.paragraphs[0].add_run(title_love + des_love)
    content_cell.paragraphs[0].font_size = shared.Inches(0.5)
    content_run.font.size = Pt(12)
    content_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    content_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    if isSend:
        content_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


def love_to_docx(conRemark, timeStart, timeEnd):
    """
    主函数，将给爱人聊天记录导出到docx

    :param conRemark:
    :param timeStart:
    :param timeEnd:
    """
    voice_love_dic = './' + timeStart[:4] + '/' + timeStart[5:7] + '/voice_love'
    file_love_dic = './' + timeStart[:4] + '/' + timeStart[5:7] + '/file_love'
    tmp_voice_dic = './' + timeStart[:4] + '/' + timeStart[5:7] + '/tmp_voice'
    video_love_dic = './' + timeStart[:4] + '/' + timeStart[5:7] + '/video_love'
    mkdir(voice_love_dic)
    mkdir(file_love_dic)
    mkdir(tmp_voice_dic)
    mkdir(video_love_dic)
    mkdir('./tmp')
    filename = f"./" + timeStart[:4] + "/" + timeStart[5:7] + "/" + timeStart[:7] + ".docx"
    doc_love = docx.Document()
    doc_love.styles['Normal'].font.name = 'Times New Roman'
    doc_love.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
    special_love_dict = get_love_msg(conRemark, timeStart, timeEnd)
    last_time = timeStart
    last_print = timeStart
    last_isSend = 520
    for love_msg in special_love_dict:
        if IS_3_min(last_time, love_msg['theTime']):
            text_love_time(doc_love, love_msg['theTime'])
        elif last_isSend ^ love_msg['isSend']:
            love_paragraph = doc_love.add_paragraph()
            love_paragraph.paragraph_format.space_before = Pt(5)
            love_paragraph.paragraph_format.space_after = Pt(5)
        last_time = love_msg['theTime']

        if IS_8_hour(last_print, love_msg['theTime']):
            print(colored(love_msg['theTime'] + ' is finished.', "green"))
            last_print = love_msg['theTime']

        if love_msg['type'] == TYPE_MSG:
            # print(love_msg['theTime'] + ' ' + love_msg['person'] + ' ' + love_msg['message'])
            text_love(doc_love, love_msg['isSend'], love_msg['person'], love_msg['message'], love_msg['status'])
        elif love_msg['type'] == TYPE_IMG:
            image_love(doc_love, love_msg['isSend'], love_msg['person'], love_msg['imgPath'])
            # print(love_msg['theTime'] + ' ' + love_msg['person'] + ' ' + love_msg['imgPath'])
        elif love_msg['type'] == TYPE_BIG_EMOJI:
            emoji_love(doc_love, love_msg['isSend'], love_msg['person'], love_msg['message'], love_msg['imgPath'])
            # print(love_msg['theTime'] + ' ' + love_msg['person'] + ' ' + love_msg['imgPath'])
        elif love_msg['type'] == TYPE_SPEAK:
            voice_love(doc_love, love_msg['isSend'], love_msg['person'], love_msg['imgPath'], tmp_voice_dic)
            # print(love_msg['theTime'] + ' ' + love_msg['person'] + ' ' + love_msg['imgPath'])
        elif love_msg['type'] == TYPE_ANSWER_MSG:
            reply_love(doc_love, love_msg['isSend'], love_msg['person'], love_msg['message'], love_msg['status'])
            # print(love_msg['theTime'] + ' ' + love_msg['person'] + ' ' + love_msg['message'])
        elif love_msg['type'] == TYPE_WITHDRAW_MSG:
            retract_message_love(doc_love, love_msg['message'])
            # print(love_msg['theTime'] + ' ' + love_msg['person'] + ' ' + love_msg['message'])
        elif love_msg['type'] == TYPE_NAME_CARD:
            namecard_love(doc_love, love_msg['isSend'], love_msg['person'], love_msg['message'])
            # print(love_msg['theTime'] + ' ' + love_msg['person'] + ' ' + love_msg['message'])
        elif love_msg['type'] == TYPE_VOIP:
            # print(love_msg['theTime'] + ' ' + love_msg['person'] + ' ' + str(love_msg['buffer'], 'utf-8'))
            voip_love(doc_love, love_msg['isSend'], love_msg['person'], love_msg['buffer'])
        elif love_msg['type'] == TYPE_FILE:
            file_love(doc_love, love_msg['isSend'], love_msg['person'], love_msg['message'], file_love_dic)
            # print(love_msg['theTime'] + ' ' + love_msg['person'] + ' ' + love_msg['message'])
        elif love_msg['type'] == TYPE_SHOT:
            shot_love(doc_love, love_msg['person'], love_msg['message'])
            # print(love_msg['theTime'] + ' ' + love_msg['person'] + ' ' + love_msg['message'])
        elif love_msg['type'] == TYPE_VIDEO_FILE:
            # print(love_msg['theTime'] + ' ' + love_msg['person'] + ' ' + love_msg['imgPath'])
            video_love(doc_love, love_msg['isSend'], love_msg['person'], love_msg['imgPath'], video_love_dic)
        elif love_msg['type'] == TYPE_SYSTEM:
            system_love(doc_love, love_msg['person'], love_msg['message'])
            # print(love_msg['theTime'] + ' ' + love_msg['person'] + ' ' + love_msg['message'])
        elif love_msg['type'] == TYPE_LINK:
            link_love(doc_love, love_msg['isSend'], love_msg['person'], love_msg['message'], file_love_dic)
            # print(love_msg['theTime'] + ' ' + love_msg['person'] + ' ' + love_msg['message'])
        elif love_msg['type'] == TYPE_CUSTOM_EMOJI:
            custom_emoji_love(doc_love, love_msg['isSend'], love_msg['person'], love_msg['message'])
            # print(love_msg['theTime'] + ' ' + love_msg['person'] + ' ' + love_msg['message'])
        elif love_msg['type'] == TYPE_MONEY_TRANSFER:
            money_transfer_love(doc_love, love_msg['isSend'], love_msg['person'], love_msg['message'])
            # print(love_msg['theTime'] + ' ' + love_msg['person'] + ' ' + love_msg['message'])
        elif love_msg['type'] == TYPE_LUCKY_MONEY:
            luck_money_love(doc_love, love_msg['isSend'], love_msg['person'], love_msg['message'])
            # print(love_msg['theTime'] + ' ' + love_msg['person'] + ' ' + love_msg['message'])
        elif love_msg['type'] == TYPE_APP_MSG:
            app_msg_love(doc_love, love_msg['isSend'], love_msg['person'], love_msg['message'])
            # print(love_msg['theTime'] + ' ' + love_msg['person'] + ' ' + love_msg['message'])
        else:
            print('wrong')
            print(love_msg['theTime'] + ' ' + love_msg['person'] + ' ' + love_msg['message'], love_msg['type'])

        last_isSend = love_msg['isSend']

    doc_love.save(filename)

    if len(os.listdir(file_love_dic)) == 0:
        os.removedirs(file_love_dic)
        # print('成功删除 ' + file_love_dic)
    if len(os.listdir(tmp_voice_dic)) == 0:
        os.removedirs(tmp_voice_dic)
        # print('成功删除 ' + tmp_voice_dic)
        os.removedirs(voice_love_dic)
        # print('成功删除 ' + voice_love_dic)
    if len(os.listdir(video_love_dic)) == 0:
        os.removedirs(video_love_dic)
        # print('成功删除 ' + video_love_dic)


def love_all(conRemark, timeStart, timeEnd):
    """
    按年月分别导出

    :param conRemark:
    :param timeStart:
    :param timeEnd:
    """
    Start_month = (int(timeStart[:4]) - 2020) * 12 + int(timeStart[5:7]) - 1
    End_month = (int(timeEnd[:4]) - 2020) * 12 + int(timeEnd[5:7]) - 1
    while Start_month < End_month:
        new_timeStart = str(Start_month // 12 + 2020) + '-' + '{:0>2d}'.format(Start_month % 12 + 1) + '-01 00:00:00'
        Start_month += 1
        new_timeEnd = str(Start_month // 12 + 2020) + '-' + '{:0>2d}'.format(Start_month % 12 + 1) + '-01 00:00:00'
        love_to_docx(conRemark, new_timeStart, new_timeEnd)
        print(colored(new_timeStart[:7] + ' is finished!!!', "red"))


def sum_love_voip(conRemark):
    """
    计算语音/视频通话时长和
    sum_love(hour, min, mes)：通话时长和/（时，分，秒）记录

    :param conRemark: 备注
    """
    select_love_voip = 'SELECT datetime(subStr(cast(msg.createTime as text),1,10),"unixepoch","localtime") AS theTime,' \
                       'msg.lvbuffer AS buffer ' \
                       'FROM message msg INNER JOIN rcontact name ON msg.talker = name.username ' \
                       'WHERE name.conRemark = \'' + conRemark + '\' ' \
                                                                 'AND msg.type = 50 ' \
                                                                 'ORDER BY theTime;'
    msg_cur, file_cur, msg_con, file_con = connect_wcdb()
    msg_cur.execute(select_love_voip)
    desc = msg_cur.description
    special_love_voip = [dict(zip([col[0] for col in desc], row)) for row in msg_cur.fetchall()]
    close_wcdb(msg_cur, file_cur, msg_con, file_con)

    sum_love = 0
    str_love = '语音通话'
    for love_voip in special_love_voip:
        str_tmp = str(love_voip['buffer'], 'utf-8')
        if str_tmp.find('聊天时长') != -1:
            str_love = str_tmp[str_tmp.find('聊天时长'):str_tmp.find(':') + 3]
        elif str_tmp.find('通话时长') != -1:
            str_love = str_tmp[str_tmp.find('通话时长'):str_tmp.find(':') + 3]
        elif str_tmp.find('通话中断') != -1:
            try:
                str_love = str_tmp[str_tmp.find('通话中断'):str_tmp.find(':') + 3]
            except Exception:
                str_love = '通话中断'
        if str_love:
            # print(str_love)
            try:
                min_love = str_love[str_love.find(' ') + 1:str_love.find(':')]
                mes_love = str_love[str_love.find(':') + 1:]
                sum_love += int(mes_love) + int(min_love) * 60
                # print(mes_love)
                # print(min_love)
            except Exception:
                sum_love = sum_love
                # print(str_love)
    minute, mes = divmod(sum_love, 60)
    hour, minute = divmod(minute, 60)
    print("%d:%02d:%02d" % (hour, minute, mes))
    print(sum_love)


def sum_love_baby(conRemark):
    """
    计算称呼宝宝的次数
    sum_love：我叫宝宝的次数
    start_love_baby：第一次叫宝宝的时间

    :param conRemark:
    """
    select_love_baby = 'SELECT datetime(subStr(cast(msg.createTime as text),1,10),"unixepoch","localtime") AS theTime,' \
                       'msg.content AS message ' \
                       'FROM message msg INNER JOIN rcontact name ON msg.talker = name.username ' \
                       'WHERE name.conRemark = \'' + conRemark + '\' ' \
                                                                 'AND msg.type = 1 ' \
                                                                 "AND msg.content LIKE '%宝宝%' " \
                                                                 'ORDER BY theTime;'
    msg_cur, file_cur, msg_con, file_con = connect_wcdb()
    msg_cur.execute(select_love_baby)
    desc = msg_cur.description
    special_love_baby = [dict(zip([col[0] for col in desc], row)) for row in msg_cur.fetchall()]
    close_wcdb(msg_cur, file_cur, msg_con, file_con)
    start_love_baby = special_love_baby[0]['theTime']
    print(start_love_baby)

    sum_love = 0

    for love_baby in special_love_baby:
        sum_love += love_baby['message'].count('宝宝')

    print(sum_love)


def sum_love_guai(conRemark):
    """
    计算称呼乖乖的次数
    sum_love：我叫乖乖的次数
    start_love_guai：第一次叫乖乖的时间

    :param conRemark:
    """
    select_love_guai = 'SELECT datetime(subStr(cast(msg.createTime as text),1,10),"unixepoch","localtime") AS theTime,' \
                       'msg.content AS message ' \
                       'FROM message msg INNER JOIN rcontact name ON msg.talker = name.username ' \
                       'WHERE name.conRemark = \'' + conRemark + '\' ' \
                                                                 'AND msg.type = 1 ' \
                                                                 "AND msg.content LIKE '%乖乖%' " \
                                                                 'ORDER BY theTime;'
    msg_cur, file_cur, msg_con, file_con = connect_wcdb()
    msg_cur.execute(select_love_guai)
    desc = msg_cur.description
    special_love_guai = [dict(zip([col[0] for col in desc], row)) for row in msg_cur.fetchall()]
    close_wcdb(msg_cur, file_cur, msg_con, file_con)
    start_love_guai = special_love_guai[0]['theTime']
    print(start_love_guai)

    sum_love = 0

    for love_guai in special_love_guai:
        sum_love += love_guai['message'].count('乖乖')

    print(sum_love)


def sum_love_ILOVEU(conRemark):
    """
    计算说我爱你的次数
    sum_love：说我爱你的次数
    start_love_ILOVEU：第一次说我爱你的时间

    :param conRemark:
    """
    select_love_ILOVEU = 'SELECT datetime(subStr(cast(msg.createTime as text),1,10),"unixepoch","localtime") AS theTime,' \
                         'msg.content AS message ' \
                         'FROM message msg INNER JOIN rcontact name ON msg.talker = name.username ' \
                         'WHERE name.conRemark = \'' + conRemark + '\' ' \
                                                                   'AND msg.type = 1 ' \
                                                                   "AND msg.content LIKE '%我爱你%' " \
                                                                   'ORDER BY theTime;'
    msg_cur, file_cur, msg_con, file_con = connect_wcdb()
    msg_cur.execute(select_love_ILOVEU)
    desc = msg_cur.description
    special_love_ILOVEU = [dict(zip([col[0] for col in desc], row)) for row in msg_cur.fetchall()]
    close_wcdb(msg_cur, file_cur, msg_con, file_con)
    start_love_ILOVEU = special_love_ILOVEU[0]['theTime']
    print(start_love_ILOVEU)

    sum_love = 0

    for love_ILOVEU in special_love_ILOVEU:
        sum_love += love_ILOVEU['message'].count('我爱你')

    print(sum_love)


def sum_love_LOVEU(conRemark):
    """
    计算说爱你的次数
    sum_love：说爱你的次数
    start_love_ILOVEU：第一次说爱你的时间

    :param conRemark:
    """
    select_love_LOVEU = 'SELECT datetime(subStr(cast(msg.createTime as text),1,10),"unixepoch","localtime") AS theTime,' \
                        'msg.content AS message ' \
                        'FROM message msg INNER JOIN rcontact name ON msg.talker = name.username ' \
                        'WHERE name.conRemark = \'' + conRemark + '\' ' \
                                                                  'AND msg.type = 1 ' \
                                                                  "AND msg.content LIKE '%爱你%' " \
                                                                  'ORDER BY theTime;'
    msg_cur, file_cur, msg_con, file_con = connect_wcdb()
    msg_cur.execute(select_love_LOVEU)
    desc = msg_cur.description
    special_love_LOVEU = [dict(zip([col[0] for col in desc], row)) for row in msg_cur.fetchall()]
    close_wcdb(msg_cur, file_cur, msg_con, file_con)
    start_love_LOVEU = special_love_LOVEU[0]['theTime']
    print(start_love_LOVEU)

    sum_love = 0

    for love_LOVEU in special_love_LOVEU:
        sum_love += love_LOVEU['message'].count('爱你')

    print(sum_love)


def sum_love_LOVE(conRemark):
    """
    计算说爱的次数
    sum_love：说爱的次数
    start_love_ILOVEU：第一次说爱的时间

    :param conRemark:
    """
    select_love_LOVE = 'SELECT datetime(subStr(cast(msg.createTime as text),1,10),"unixepoch","localtime") AS theTime,' \
                       'msg.content AS message ' \
                       'FROM message msg INNER JOIN rcontact name ON msg.talker = name.username ' \
                       'WHERE name.conRemark = \'' + conRemark + '\' ' \
                                                                 'AND msg.type = 1 ' \
                                                                 "AND msg.content LIKE '%爱%' " \
                                                                 'ORDER BY theTime;'
    msg_cur, file_cur, msg_con, file_con = connect_wcdb()
    msg_cur.execute(select_love_LOVE)
    desc = msg_cur.description
    special_love_LOVE = [dict(zip([col[0] for col in desc], row)) for row in msg_cur.fetchall()]
    close_wcdb(msg_cur, file_cur, msg_con, file_con)
    start_love_LOVE = special_love_LOVE[0]['theTime']
    print(start_love_LOVE)

    sum_love = 0

    for love_LOVE in special_love_LOVE:
        sum_love += love_LOVE['message'].count('爱')

    print(sum_love)


def love_in_night(conRemark):
    """
    记录每一次深夜畅聊
    sum_love_talk：两人凌晨还在聊天的天数
    sum_love：至少有人熬夜到第二天发消息的天数
    trigger_love：发消息卡凌晨00：00的天数
    start_love_in_night：第一次因为对方熬夜到第二天

    :param conRemark:
    """
    select_love_talk = 'SELECT datetime(subStr(cast(msg.createTime as text),1,10),"unixepoch","localtime") AS theTime ' \
                       'FROM message msg INNER JOIN rcontact name ON msg.talker = name.username ' \
                       'WHERE name.conRemark = \'' + conRemark + '\' ' \
                                                                 'AND (msg.type = 50 OR msg.isSend = 0) ' \
                                                                 'ORDER BY theTime;'
    msg_cur, file_cur, msg_con, file_con = connect_wcdb()
    msg_cur.execute(select_love_talk)
    desc = msg_cur.description
    special_love_talk = [dict(zip([col[0] for col in desc], row)) for row in msg_cur.fetchall()]
    close_wcdb(msg_cur, file_cur, msg_con, file_con)

    timestamp_flag = 23400

    love_flag = "00"
    sum_love_talk = 0

    for love_talk in special_love_talk:
        if love_talk['theTime'][8:10] == love_flag:
            continue
        love_flag = love_talk['theTime'][8:10]
        love_timestamp = int(love_talk['theTime'][11:13]) * 60 * 60 + int(love_talk['theTime'][14:16]) * 60 + int(
            love_talk['theTime'][17:19])
        if love_timestamp < timestamp_flag:
            sum_love_talk += 1

    select_love = 'SELECT datetime(subStr(cast(msg.createTime as text),1,10),"unixepoch","localtime") AS theTime ' \
                  'FROM message msg INNER JOIN rcontact name ON msg.talker = name.username ' \
                  'WHERE name.conRemark = \'' + conRemark + '\' ' \
                                                            'ORDER BY theTime;'
    msg_cur, file_cur, msg_con, file_con = connect_wcdb()
    msg_cur.execute(select_love)
    desc = msg_cur.description
    special_love = [dict(zip([col[0] for col in desc], row)) for row in msg_cur.fetchall()]
    close_wcdb(msg_cur, file_cur, msg_con, file_con)

    timestamp_flag = 23400

    trigger_love = 0
    love_flag = "00"
    sum_love = 0
    flag = 1
    for love in special_love:
        if love['theTime'][8:10] == love_flag:
            continue
        love_flag = love['theTime'][8:10]
        love_timestamp = int(love['theTime'][11:13]) * 60 * 60 + int(love['theTime'][14:16]) * 60 + int(
            love['theTime'][17:19])
        if love_timestamp < timestamp_flag:
            if flag == 1:
                start_love_in_night = love['theTime']
                flag = 0
            sum_love += 1
        if love['theTime'][11:16] == '00:00':
            trigger_love += 1

    print(sum_love_talk)
    print(sum_love)
    print(trigger_love)
    print(start_love_in_night)


def sum_love_wanan(conRemark):
    """
    计算说晚安的次数
    sum_love：说晚安的次数
    start_love_wanan：第一次说晚安的时间

    :param conRemark:
    """
    select_love_wanan = 'SELECT datetime(subStr(cast(msg.createTime as text),1,10),"unixepoch","localtime") AS theTime,' \
                        'msg.content AS message ' \
                        'FROM message msg INNER JOIN rcontact name ON msg.talker = name.username ' \
                        'WHERE name.conRemark = \'' + conRemark + '\' ' \
                                                                  'AND msg.type = 1 ' \
                                                                  "AND msg.content LIKE '%晚安%' " \
                                                                  'ORDER BY theTime;'
    msg_cur, file_cur, msg_con, file_con = connect_wcdb()
    msg_cur.execute(select_love_wanan)
    desc = msg_cur.description
    special_love_wanan = [dict(zip([col[0] for col in desc], row)) for row in msg_cur.fetchall()]
    close_wcdb(msg_cur, file_cur, msg_con, file_con)
    start_love_wanan = special_love_wanan[0]['theTime']
    print(start_love_wanan)

    sum_love = 0

    for love_wanan in special_love_wanan:
        sum_love += love_wanan['message'].count('晚安')

    print(sum_love)


def get_love_sum(conRemark, timeStart, timeInterval):
    """
    记录当前时间间隔发送消息数量

    :param conRemark: 备注
    :param timeStart: 开始时间（时间戳表示）
    :param timeInterval: 时间间隔长度（时间戳表示）
    :return: timeEnd, love_sum: 结束时间， 消息数量
    """
    timeEnd = timeStart + timeInterval
    select_love_sum = 'SELECT COUNT(*) ' \
                      'FROM message msg INNER JOIN rcontact name ON msg.talker = name.username ' \
                      'WHERE name.conRemark = \'' + conRemark + '\' ' \
                                                                'AND msg.createTime < ' + str(timeEnd) + ' ' \
                                                                'AND msg.createTime >= ' + str(timeStart) + ' ' \
                                                                'ORDER BY msg.createTime;'
    msg_cur, file_cur, msg_con, file_con = connect_wcdb()
    msg_cur.execute(select_love_sum)
    love_sum = msg_cur.fetchone()[0]
    close_wcdb(msg_cur, file_cur, msg_con, file_con)
    return timeEnd, love_sum


def sum_love_every(conRemark, Start, End, Interval):
    """
    绘制每天聊天记录曲线图，且记载次数最多的时间和数量，以及全部总和

    :param conRemark:
    :param Start:
    :param End:
    :param Interval: 时间间隔长度（/天）
    """
    timeStart = int(time.mktime(time.strptime(Start, '%Y-%m-%d %H:%M:%S'))) * 1000
    timeEnd = int(time.mktime(time.strptime(End, '%Y-%m-%d %H:%M:%S'))) * 1000
    timeInterval = 86400000 * Interval
    love_list = []
    love_time = []
    max_love = 0
    max_love_time = Start[:10]
    while timeStart < timeEnd:
        love_time.append(time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(timeStart / 1000))[:10])
        timeStart, love_sum = get_love_sum(conRemark, timeStart, timeInterval)
        love_list.append(love_sum)
        if love_sum >= max_love:
            max_love = love_sum
            max_love_time = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(timeStart / 1000))[:10]
    fig, ax = plt.subplots(figsize=(7, 3), dpi=200)

    ax.spines["left"].set_visible(False)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    ax.grid(ls="--", lw=0.5, color="#4E616C")
    print(max_love)
    print(max_love_time)
    print(sum(love_list))
    ax.plot(love_time, love_list, marker="o", mfc="white", ms=2, lw=1)
    ax.xaxis.set_major_locator(ticker.MultipleLocator(20))  # ticker every 2 matchdays
    ax.set_xticks(love_time[::20])
    ax.xaxis.set_tick_params(direction='inout', length=2, color="#4E616C", labelcolor="#4E616C", labelsize=4,
                             rotation=90)
    ax.yaxis.set_tick_params(direction='inout', length=2, color="#4E616C", labelcolor="#4E616C", labelsize=5)

    ax.spines["bottom"].set_edgecolor("#4E616C")
    plt.show()
